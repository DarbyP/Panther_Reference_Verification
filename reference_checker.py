"""
Panther Reference Verification
A tool to detect potentially fabricated references in student papers.

Checks references against CrossRef and PubMed (journals), Open Library and Google Books (books).

⚠️ IMPORTANT: WHAT THIS TOOL DOES AND DOES NOT DO
===================================================

This tool is a SCREENING AID that helps instructors identify references that may need 
manual verification. It does NOT definitively prove that a reference is falsified.

✓ WHAT IT DOES:
- Automatically verifies references against major academic databases
- Filters out references that are found and verified
- Flags references that could not be verified for manual review
- Reduces the number of references instructors need to manually check

✗ WHAT IT DOES NOT DO:
- Does NOT prove a reference is fabricated (many legitimate reasons for flags)
- Does NOT replace instructor judgment and manual verification
- Does NOT check all sources (some journals/books are not indexed)

⚠️ ALL FLAGGED REFERENCES REQUIRE MANUAL VERIFICATION BY THE INSTRUCTOR

Common reasons for legitimate flags:
• Journal not indexed in CrossRef/PubMed
• Book not in Open Library/Google Books
• Student made an honest transcription error (misspelled title, wrong year)
• Non-English sources
• Very recent publications not yet indexed
• Regional or specialized publications

The goal is to REDUCE verification burden, not eliminate it entirely.

PROCESSING FLOW:
================

1. STYLE DETECTION (detect_citation_style)
   - Analyzes reference list and in-text citations
   - Returns: 'APA', 'MLA', or 'UNKNOWN' with confidence %
   - This determines ALL subsequent processing

2. REFERENCE EXTRACTION & SPLITTING
   APA Style Guide:
   - Format: Author, I. N. (Year). Title. Source.
   - Split by: Year in parentheses at start
   - Function: split_apa_references()
   
   MLA Style Guide:
   - Format: Author, Full Name. "Title." Container, elements, year.
   - Split by: "LastName, FirstName." or quoted title
   - Filters: Page headers ("LastName PageNumber")
   - Function: split_mla_references()

3. REFERENCE PARSING
   APA Style Guide:
   - Author: Everything before (Year)
   - Year: In parentheses early
   - Title: After year, before source
   - Function: parse_apa_reference()
   
   MLA Style Guide:
   - Author: Before first period (if not starting with quote)
   - Title: Text in quotes
   - Year: Last 50 chars (end of reference)
   - Container: Between title and year
   - Function: parse_mla_reference()

4. CITATION EXTRACTION
   APA Style Guide:
   - Parenthetical: (Author, Year)
   - Narrative: Author (Year)
   - Function: extract_intext_citations() with citation_style='APA'
   
   MLA Style Guide:
   - Parenthetical: (Author Page) or (Author)
   - Narrative: Author names from ref list in scholarly contexts
   - Function: extract_intext_citations() with citation_style='MLA'

5. CITATION-REFERENCE MATCHING
   - Normalizes author names (preserves order for et al.)
   - Handles MLA (no year) vs APA (requires year)
   - Function: match_citations_to_references()

6. VERIFICATION
   - Checks references against academic databases
   - Style-agnostic (works on parsed components)
   - Function: verify_all_references()
"""

import os
import sys
import re
import uuid
import time
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, Menu
from tkinter import PhotoImage
from pathlib import Path
from datetime import datetime
from collections import defaultdict
import webbrowser

import requests
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pdfplumber
from PIL import Image, ImageTk

# =============================================================================
# VERSION AND UPDATE CONFIGURATION
# =============================================================================
VERSION = "1.0.2"
GITHUB_REPO = "DarbyP/Panther_Reference_Verification"  # Update this with your GitHub username
GITHUB_API_RELEASES = f"https://api.github.com/repos/{GITHUB_REPO}/releases/latest"

# Colors
FT_CRIMSON = '#770000'
FT_CRIMSON_RGB = RGBColor(0x77, 0x00, 0x00)
FT_ORANGE_RGB = RGBColor(0xCC, 0x66, 0x00)
FT_GREEN_RGB = RGBColor(0x00, 0x66, 0x00)
FT_WHITE = '#FFFFFF'
FT_LIGHT_GRAY = '#F5F5F5'

# API settings
CROSSREF_API = "https://api.crossref.org/works"
HEADERS = {'User-Agent': 'PantherReferenceVerification/1.0 (Academic integrity tool)'}


def get_resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller bundle."""
    if hasattr(sys, '_MEIPASS'):
        # Running in a PyInstaller bundle
        base_path = sys._MEIPASS
    else:
        # Running in development
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)


def check_for_updates():
    """Check GitHub releases for a newer version. Returns (has_update, latest_version, download_url) or None on error."""
    try:
        response = requests.get(GITHUB_API_RELEASES, timeout=5)
        if response.status_code == 200:
            data = response.json()
            latest_version = data.get('tag_name', '').lstrip('v')
            download_url = data.get('html_url', '')
            
            # Compare versions
            def parse_version(v):
                return tuple(int(x) for x in v.split('.'))
            
            try:
                if parse_version(latest_version) > parse_version(VERSION):
                    return (True, latest_version, download_url)
            except:
                pass
            return (False, latest_version, download_url)
    except:
        pass
    return None


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    # Get the document part
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create run element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    # Add font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')  # 9pt = 18 half-points
    rPr.append(sz)
    
    new_run.append(rPr)
    
    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


# =============================================================================
# COMPONENT 1: File Ingestion & Student Identification
# =============================================================================

def extract_text_from_docx(filepath):
    """Extract all text from a DOCX file."""
    doc = DocxDocument(filepath)
    return [p.text.strip() for p in doc.paragraphs]


def extract_text_from_pdf(filepath):
    """Extract all text from a PDF file, split by lines."""
    paragraphs = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                paragraphs.extend(text.split('\n'))
    return [p.strip() for p in paragraphs]


def extract_text_from_pdf_page(filepath, page_number):
    """Extract text from a specific page of a PDF file (1-indexed)."""
    paragraphs = []
    with pdfplumber.open(filepath) as pdf:
        if 0 < page_number <= len(pdf.pages):
            page = pdf.pages[page_number - 1]  # Convert to 0-indexed
            text = page.extract_text()
            if text:
                paragraphs.extend(text.split('\n'))
    return [p.strip() for p in paragraphs]


def extract_student_name(paragraphs):
    """Extract student name from APA title page."""
    first_lines = [p for p in paragraphs if p][:15]
    
    for line in first_lines:
        if len(line) < 3 or len(line) > 50:
            continue
        
        lower = line.lower()
        skip_keywords = [
            'university', 'college', 'department', 'course', 'professor',
            'instructor', 'dr.', 'dr ', 'january', 'february', 'march',
            'april', 'may', 'june', 'july', 'august', 'september',
            'october', 'november', 'december', 'running head', 'abstract',
            'introduction', 'assignment', 'paper', 'final', 'midterm',
            'psyc', 'psych', '101', '201', '301', '401'
        ]
        if any(kw in lower for kw in skip_keywords):
            continue
        
        words = line.split()
        if 2 <= len(words) <= 4:
            alpha_chars = sum(c.isalpha() for c in line)
            if alpha_chars / max(len(line.replace(' ', '')), 1) > 0.9:
                return line
    
    return "Unknown"


def ingest_papers(folder_path):
    """Process all DOCX and PDF files in a folder."""
    folder = Path(folder_path)
    lookup = {}
    
    files = list(folder.glob('*.docx')) + list(folder.glob('*.pdf'))
    files += list(folder.glob('*.DOCX')) + list(folder.glob('*.PDF'))
    
    for filepath in files:
        if filepath.name.startswith('~$'):
            continue
        
        suffix = filepath.suffix.lower()
        if suffix == '.docx':
            paragraphs = extract_text_from_docx(filepath)
            student_name = extract_student_name(paragraphs)
        elif suffix == '.pdf':
            # Check if this is a Turnitin PDF (has "Submission ID" on first page)
            first_page_text = extract_text_from_pdf_page(filepath, 1)
            is_turnitin = any('Submission ID' in line for line in first_page_text)
            
            if is_turnitin:
                # For Turnitin PDFs, extract student name from page 3
                page_3_text = extract_text_from_pdf_page(filepath, 3)
                student_name = extract_student_name(page_3_text)
                # Still get all paragraphs for reference extraction
                paragraphs = extract_text_from_pdf(filepath)
            else:
                # Normal PDF processing
                paragraphs = extract_text_from_pdf(filepath)
                student_name = extract_student_name(paragraphs)
        else:
            continue
        
        code = f"REF_{uuid.uuid4().hex[:6].upper()}"
        
        lookup[code] = {
            'student_name': student_name,
            'filepath': str(filepath),
            'paragraphs': paragraphs
        }
    
    return lookup


# =============================================================================
# COMPONENT 1.5: Citation Style Detection  
# =============================================================================

def detect_citation_style(paragraphs):
    """
    Detect whether a paper uses APA or MLA citation style.
    
    Strategy: Analyze the reference list format (most reliable indicator)
    - APA: Author, I. N. (Year). Title.
    - MLA: Author, Full Name. "Title."
    
    Returns: (style, confidence_percentage)
    """
    # Find the references section
    ref_section = find_references_section(paragraphs)
    
    if not ref_section or len(ref_section) < 2:
        return 'UNKNOWN', 0
    
    score = {'APA': 0, 'MLA': 0}
    
    # Check header (if we can find it)
    for i, p in enumerate(paragraphs):
        cleaned = p.strip().lower()
        if re.match(r'^references?\s*$', cleaned):
            score['APA'] += 2
            break
        elif re.match(r'^works\s+cited\s*$', cleaned):
            score['MLA'] += 3  # Strong signal
            break
    
    # Analyze first 5 reference entries (or fewer if less available)
    sample_refs = ref_section[:min(5, len(ref_section))]
    
    for ref in sample_refs:
        if len(ref) < 20:  # Skip very short lines
            continue
        
        # APA signatures:
        # 1. Initials after last name: "Smith, J. A."
        # 2. Year in parentheses early in entry: "(2020)"
        # 3. Pattern: Author, X. Y. (YEAR).
        
        # Look for initials pattern (one or two capital letters with periods)
        if re.search(r'[A-Z][a-z]+,\s+[A-Z]\.\s*(?:[A-Z]\.\s*)?\(', ref):
            score['APA'] += 3
        
        # Look for year in parentheses near the start
        if re.search(r'^[^(]{10,60}\(\d{4}[a-z]?\)', ref):
            score['APA'] += 2
        
        # MLA signatures:
        # 1. Full first name: "Smith, John" or "Smith, John Andrew"
        # 2. Title in quotes or italics immediately after author
        # 3. Year appears later, often near end, no parentheses
        
        # Look for full name pattern (capital first name after comma)
        if re.search(r'[A-Z][a-z]+,\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\.\s+"', ref):
            score['MLA'] += 3
        
        # Look for full first name without immediate parentheses
        if re.search(r'[A-Z][a-z]+,\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\.', ref) and not re.search(r'\(\d{4}\)', ref[:60]):
            score['MLA'] += 2
        
        # Look for title in quotes early in entry (MLA pattern)
        if re.search(r'^[^"]{10,50}"[^"]{10,}".{20,}\d{4}', ref):
            score['MLA'] += 2
    
    # Also check in-text citation patterns in body (secondary signal)
    body = extract_paper_body(paragraphs)
    body_text = ' '.join(body[:100])  # Sample first 100 paragraphs
    
    # APA: (Author, Year) with comma
    apa_citations = len(re.findall(r'\([A-Z][a-z]+(?:\s+et al\.)?,\s*\d{4}\)', body_text))
    
    # MLA: (Author Page) or just (Author) - space, no comma before number
    mla_citations = len(re.findall(r'\([A-Z][a-z]+\s+\d{1,4}(?:-\d{1,4})?\)', body_text))
    
    if apa_citations > mla_citations * 2:  # Strong APA signal
        score['APA'] += 2
    elif mla_citations > apa_citations:
        score['MLA'] += 2
    
    # Calculate confidence
    total_score = score['APA'] + score['MLA']
    
    if total_score == 0:
        return 'UNKNOWN', 0
    
    if score['APA'] > score['MLA']:
        confidence = (score['APA'] / total_score) * 100
        return 'APA', confidence
    elif score['MLA'] > score['APA']:
        confidence = (score['MLA'] / total_score) * 100
        return 'MLA', confidence
    else:
        return 'UNKNOWN', 50


# =============================================================================
# COMPONENT 2: Reference Section Extraction
# =============================================================================

def clean_turnitin_footer(text):
    """
    Remove Turnitin footer patterns from reference text.
    
    Turnitin adds footers like:
    - "Page 12 of 12 - AI Writing Submission Submission ID trn:oid:::1:3427731680"
    - "Page 5 of 12"
    - Other similar patterns with Submission ID
    
    This function is more permissive to catch footers even when spacing is inconsistent.
    """
    # Pattern 1: "Page X of Y" followed by anything (most common)
    # Use \s* instead of \s+ to allow zero or more spaces before "Page"
    # This catches cases like "file.pdf Page 12" or "file.pdfPage 12" (no space)
    text = re.sub(r'\s*Page\s+\d+\s+of\s+\d+.*$', '', text, flags=re.IGNORECASE)
    
    # Pattern 2: Also remove if "Page X of Y" appears in the middle (not just at end)
    # This handles cases where multiple references are on one line
    text = re.sub(r'\s*Page\s+\d+\s+of\s+\d+[^\w]*(?:AI\s+Writing|Submission).*?(?=\s[A-Z][a-z]+,|\s*$)', '', text, flags=re.IGNORECASE)
    
    # Pattern 3: "Submission ID" patterns if they somehow appear alone
    text = re.sub(r'\s*Submission\s+ID\s+[^\s]+.*$', '', text, flags=re.IGNORECASE)
    
    # Pattern 4: "AI Writing Submission" if it somehow appears alone
    text = re.sub(r'\s*AI\s+Writing\s+Submission.*$', '', text, flags=re.IGNORECASE)
    
    # Pattern 5: Standalone "Page X of Y -" or "Page X of Y" at the end
    # More aggressive - removes even partial footer patterns
    text = re.sub(r'\s*Page\s+\d+\s+of\s+\d+\s*-?\s*$', '', text, flags=re.IGNORECASE)
    
    return text.strip()


def find_references_section(paragraphs):
    """Find and extract the references section, filtering out Turnitin footers."""
    start_idx = None
    for i, p in enumerate(paragraphs):
        cleaned = p.strip().lower()
        # Check for both "References" (APA) and "Works Cited" (MLA)
        if re.match(r'^references?\s*$', cleaned) or re.match(r'^works\s+cited\s*$', cleaned):
            start_idx = i + 1
            break
    
    if start_idx is None:
        return []
    
    end_idx = len(paragraphs)
    for i in range(start_idx, len(paragraphs)):
        cleaned = paragraphs[i].strip().lower()
        if re.match(r'^appendix|^appendices|^figure\s*\d|^table\s*\d', cleaned):
            end_idx = i
            break
    
    # Extract and clean each paragraph, filtering out Turnitin footers
    cleaned_refs = []
    for p in paragraphs[start_idx:end_idx]:
        if p.strip():
            cleaned = clean_turnitin_footer(p)
            if cleaned:  # Only add if there's content left after cleaning
                cleaned_refs.append(cleaned)
    
    return cleaned_refs


def extract_references_from_lookup(lookup):
    """Extract references section from each paper."""
    results = {}
    for code, info in lookup.items():
        ref_section = find_references_section(info['paragraphs'])
        
        # Detect citation style for this paper
        style, confidence = detect_citation_style(info['paragraphs'])
        
        results[code] = {
            'student_name': info['student_name'],
            'filepath': info['filepath'],
            'references_text': ref_section,
            'citation_style': style,
            'style_confidence': confidence
        }
    return results


# =============================================================================
# COMPONENT 2.5: Citation-Reference Matching
# =============================================================================

def extract_paper_body(paragraphs):
    """Extract the paper body (everything before the References section)."""
    ref_idx = None
    for i, p in enumerate(paragraphs):
        cleaned = p.strip().lower()
        if re.match(r'^references?\s*$', cleaned) or re.match(r'^works\s+cited\s*$', cleaned):
            ref_idx = i
            break
    
    if ref_idx is None:
        # No references section found, return all paragraphs
        return paragraphs
    
    # Skip title page - try to find where main content starts
    start_idx = 0
    for i in range(min(20, ref_idx)):
        p = paragraphs[i].strip().lower()
        if p in ['abstract', 'introduction']:
            start_idx = i
            break
    
    return [p for p in paragraphs[start_idx:ref_idx] if p.strip()]


def get_reference_last_names(references):
    """
    Extract a set of last names from the reference list.
    Used for MLA narrative citation detection.
    
    In MLA format, the last name is the first element before the comma:
    - "Ball, Lucille" → "Ball"
    - "Smith, John" → "Smith"
    - "Pontikes, Elizabeth, et al." → "Pontikes"
    """
    last_names = set()
    for ref in references:
        authors = ref.get('authors', '')
        if authors:
            # In MLA, the last name is everything before the first comma
            # Handle multiple authors separated by commas after the first author
            first_author = authors.split(',')[0].strip()
            
            # Clean up any remaining punctuation or "and"
            first_author = re.sub(r'\s+and\s+.*', '', first_author, flags=re.IGNORECASE)
            first_author = first_author.strip()
            
            # Take the first word (the actual last name)
            words = first_author.split()
            if words:
                last_names.add(words[0])
    
    return last_names


def extract_intext_citations(paper_body, citation_style='APA', reference_last_names=None):
    """
    Extract in-text citations from the paper body.
    Returns a list of (author, year) tuples.
    
    Handles both APA and MLA citation styles.
    
    For APA:
    - Parenthetical: (Smith, 2020)
    - Narrative: Smith (2020)
    - Multiple authors: (Smith & Jones, 2020) or (Smith et al., 2020)
    - Multiple citations: (Smith, 2020; Jones, 2019)
    - Prefixes: (e.g., Smith, 2020), (see: Smith, 2020), (cf. Jones, 2019)
    - Organizations: (American Psychological Association, 2017)
    
    For MLA:
    - Parenthetical: (Smith 45) or (Smith)
    - Narrative: mentions of author names from reference list
    """
    citations = []
    text = ' '.join(paper_body)
    
    if citation_style == 'APA':
        # Pattern for parenthetical citations with OPTIONAL prefixes
        # Now handles organization names including lowercase connector words (of, the, for, etc.)
        # Author pattern: Starts with capital word, then optionally more capital words or lowercase connectors
        author_pattern = r'[A-Z][A-Za-z\-\']+(?:\s+(?:of|the|for|and|in|on|at|to|a|an|from)\s+[A-Z][A-Za-z\-\']+|\s+[A-Z][A-Za-z\-\']+)*'
        paren_pattern = r'\((?:(?:e\.g\.,|cf\.|see|see e\.g\.,|see for example,?|see also):?\s*)?(' + author_pattern + r'(?:\s+et al\.|\s+(?:and|&)\s+(?:colleagues|co-workers)|\s+(?:&|and)\s+' + author_pattern + r')?),?\s+((?:19|20)\d{2}[a-z]?)\)'
        
        for match in re.finditer(paren_pattern, text):
            author = match.group(1)
            year = match.group(2)
            citations.append((author, year))
        
        # Pattern for multiple citations in one parenthesis
        multi_paren_pattern = r'\(([^)]+)\)'
        for match in re.finditer(multi_paren_pattern, text):
            content = match.group(1)
            # Only process if it contains semicolons (multiple citations)
            if ';' in content:
                # Remove common prefixes first
                content = re.sub(r'^(?:e\.g\.,|cf\.|see|see e\.g\.,|see for example,?|see also):?\s*', '', content, flags=re.IGNORECASE)
                parts = content.split(';')
                for part in parts:
                    # Try to find Author, Year pattern in each part
                    sub_pattern = r'(' + author_pattern + r'(?:\s+et al\.|\s+(?:and|&)\s+(?:colleagues|co-workers)|\s+(?:&|and)\s+' + author_pattern + r')?),?\s+((?:19|20)\d{2}[a-z]?)'
                    sub_match = re.search(sub_pattern, part)
                    if sub_match:
                        author = sub_match.group(1)
                        year = sub_match.group(2)
                        citations.append((author, year))
        
        # Pattern for narrative citations with organization name support
        narrative_pattern = r'(' + author_pattern + r'(?:\s+et al\.|\s+(?:and|&)\s+(?:colleagues|co-workers)|\s+(?:&|and)\s+' + author_pattern + r')?)\s+\(((?:19|20)\d{2}[a-z]?)\)'
        
        for match in re.finditer(narrative_pattern, text):
            author = match.group(1).strip()
            year = match.group(2)
            # Avoid duplicates from patterns that might overlap
            if (author, year) not in citations:
                citations.append((author, year))
    
    elif citation_style == 'MLA':
        # MLA parenthetical: (Author Page) or (Author)
        author_pattern = r'[A-Z][A-Za-z\-\']+'
        
        # Pattern 1: (Author Page)
        paren_with_page = r'\((' + author_pattern + r'(?:\s+et al\.)?)\s+(\d{1,4}(?:-\d{1,4})?)\)'
        for match in re.finditer(paren_with_page, text):
            author = match.group(1)
            # In MLA, we don't have year in citation, set to None
            citations.append((author, None))
        
        # Pattern 2: (Author) - no page
        paren_no_page = r'\((' + author_pattern + r'(?:\s+et al\.)?)\)'
        for match in re.finditer(paren_no_page, text):
            author = match.group(1)
            # Avoid double-counting if already captured with page
            if not any(c[0] == author for c in citations):
                citations.append((author, None))
        
        # Pattern 3: Narrative citations - author names from reference list appearing in text
        if reference_last_names:
            for last_name in reference_last_names:
                # Common narrative patterns in MLA
                patterns = [
                    rf'\b{last_name}\s+(?:argues|states|suggests|notes|writes|demonstrates|shows|finds|claims|asserts|observes|explains|contends|believes|maintains)',
                    rf'\b{last_name}\'s\s+(?:study|research|work|article|book|essay|analysis|findings|argument|theory)',
                    rf'(?:according to|as)\s+{last_name}\b',
                ]
                
                for pattern in patterns:
                    matches = re.finditer(pattern, text, re.IGNORECASE)
                    for match in matches:
                        # Try to find a year nearby (within 50 chars) for better matching
                        context = text[max(0, match.start()-50):match.end()+50]
                        year_match = re.search(r'\b(19|20)\d{2}\b', context)
                        year = year_match.group(0) if year_match else None
                        
                        # Avoid duplicates
                        if not any(c[0].lower() == last_name.lower() for c in citations):
                            citations.append((last_name, year))
                            break  # Only add once per author
    
    return citations


def extract_reference_authors_years(references):
    """
    Extract author names and years from parsed references.
    Returns a list of (authors_text, year) tuples.
    
    For MLA: Includes references even if year is missing (year will be empty string).
    For APA: Requires both author and year.
    """
    ref_citations = []
    for ref in references:
        authors = ref.get('authors', '')
        year = ref.get('year', '')
        # Include if we have authors (year optional for MLA)
        if authors:
            ref_citations.append((authors, year if year else ''))
    return ref_citations


def normalize_author(author_text):
    """
    Normalize author names from reference list.
    Extracts only last names, ignoring initials and punctuation.
    PRESERVES ORDER - first author matters for et al. matching!
    
    Examples:
    - "Dommeyer, C. J., Baum, P., & Hanna, R. W" -> "dommeyer baum hanna"
    - "Smith & Jones" -> "smith jones" (NOT sorted - order matters!)
    """
    # Remove any text in parentheses (like editor designations)
    text = re.sub(r'\([^)]*\)', '', author_text)
    
    # Remove "et al." completely
    text = re.sub(r'\s+et\s+al\.?', '', text, flags=re.IGNORECASE)
    
    # Split by common delimiters: commas, ampersands, "and"
    # This gives us individual author segments
    text = re.sub(r'\s+&\s+|\s+and\s+', ',', text, flags=re.IGNORECASE)
    segments = text.split(',')
    
    last_names = []
    for segment in segments:
        segment = segment.strip()
        if not segment:
            continue
        
        # Remove all initials (single capital letter optionally followed by period)
        # Pattern: space/comma/start followed by single capital letter followed by period/space/comma/end
        segment = re.sub(r'(?:^|[\s,])[A-Z]\.?(?=[\s,]|$)', ' ', segment)
        
        # Remove any remaining punctuation
        segment = re.sub(r'[,.\']', '', segment)
        
        # Clean up extra spaces
        segment = ' '.join(segment.split())
        
        # What's left should be last name(s)
        words = segment.split()
        for word in words:
            if word and len(word) > 1:
                last_names.append(word.lower())
    
    # DO NOT SORT - order matters for et al. matching!
    return ' '.join(last_names)


def extract_last_name_from_citation(citation_author):
    """
    Extract last name(s) from an in-text citation.
    PRESERVES ORDER - important for matching!
    
    Handles: "Berk", "Smith & Jones", "Smith et al.", "Smith and colleagues"
    Returns: "berk", "smith jones", "smith", "smith"
    """
    # Remove "et al.", "colleagues", and "co-workers"
    text = re.sub(r'\s+et\s+al\.?', '', citation_author, flags=re.IGNORECASE)
    text = re.sub(r'\s+(?:and|&)\s+(?:colleagues|co-workers)', '', text, flags=re.IGNORECASE)
    
    # Split by & or "and"
    text = re.sub(r'\s+&\s+|\s+and\s+', ',', text, flags=re.IGNORECASE)
    names = [n.strip() for n in text.split(',') if n.strip()]
    
    # Convert to lowercase but DO NOT SORT - order matters!
    return ' '.join([n.lower() for n in names if n])


def normalize_year(year):
    """
    Normalize year by removing letter suffixes (2015a -> 2015).
    Returns None for empty/missing years (MLA style).
    """
    if year and str(year).strip():
        return re.sub(r'[a-z]$', '', str(year))
    return None


def match_citations_to_references(intext_citations, reference_citations):
    """
    Match in-text citations to references.
    Returns:
    - uncited_refs: References that are not cited in the text
    - missing_refs: In-text citations that have no matching reference
    """
    # Create a lookup for references by normalized author and year
    # Key: (normalized_author, normalized_year) -> Value: list of original (author, year) tuples
    ref_lookup = defaultdict(list)
    
    for authors, year in reference_citations:
        # Normalize the reference author (removes initials, extracts last names)
        norm_author = normalize_author(authors)
        norm_year = normalize_year(year)
        
        # Also store just the first last name for et al. matching
        first_last_name = norm_author.split()[0] if norm_author else ''
        
        # Store both full author list and first author
        ref_lookup[(norm_author, norm_year)].append((authors, year))
        if first_last_name and norm_author != first_last_name:
            ref_lookup[(first_last_name, norm_year)].append((authors, year))
    
    # Track which references were cited (use original author strings)
    cited_refs = set()
    missing_refs = []
    
    # Check each in-text citation
    for cite_author, cite_year in intext_citations:
        # Normalize the citation author (already just last names)
        norm_cite = extract_last_name_from_citation(cite_author)
        norm_year = normalize_year(cite_year) if cite_year else None
        first_last_name = norm_cite.split()[0] if norm_cite else ''
        
        found = False
        
        # For MLA (no year in citation), try to match just by author
        if norm_year is None:
            # Look for any reference with this author
            for (ref_author, ref_year), original_refs in ref_lookup.items():
                if norm_cite == ref_author or norm_cite == ref_author.split()[0]:
                    for orig_auth, orig_year in original_refs:
                        cited_refs.add((orig_auth, orig_year))
                    found = True
                    break
        else:
            # APA style - match by author and year
            # Try exact match first (all authors)
            if (norm_cite, norm_year) in ref_lookup:
                # Mark all matching references as cited
                for ref_authors, ref_year in ref_lookup[(norm_cite, norm_year)]:
                    cited_refs.add((ref_authors, ref_year))
                found = True
            # Try first author only (for et al.)
            elif (first_last_name, norm_year) in ref_lookup:
                for ref_authors, ref_year in ref_lookup[(first_last_name, norm_year)]:
                    cited_refs.add((ref_authors, ref_year))
                found = True
        
        if not found:
            # No match found
            missing_refs.append((cite_author, cite_year))
    
    # Find uncited references (use original strings, not normalized)
    uncited_refs = []
    for authors, year in reference_citations:
        if (authors, year) not in cited_refs:
            uncited_refs.append((authors, year))
    
    # Remove duplicates from missing_refs while preserving order
    seen = set()
    unique_missing = []
    for item in missing_refs:
        if item not in seen:
            seen.add(item)
            unique_missing.append(item)
    
    return uncited_refs, unique_missing


def check_citation_matching(lookup):
    """
    Check citation-reference matching for all papers.
    Returns results with citation matching information added.
    """
    results = {}
    
    for code, info in lookup.items():
        paragraphs = info['paragraphs']
        
        # Detect citation style
        style, confidence = detect_citation_style(paragraphs)
        
        # Extract paper body and references section
        paper_body = extract_paper_body(paragraphs)
        ref_section = find_references_section(paragraphs)
        
        if not ref_section:
            results[code] = {
                'uncited_refs': [],
                'missing_refs': [],
                'total_citations': 0,
                'total_references': 0,
                'has_issues': False
            }
            continue
        
        # Split and parse references
        individual_refs = split_references(ref_section, citation_style=style)
        parsed_refs = [parse_reference(ref, citation_style=style) for ref in individual_refs]
        
        # Get reference last names for MLA narrative detection
        reference_last_names = get_reference_last_names(parsed_refs)
        
        # Extract in-text citations with style-appropriate detection
        intext_citations = extract_intext_citations(paper_body, citation_style=style,
                                                     reference_last_names=reference_last_names)
        
        reference_citations = extract_reference_authors_years(parsed_refs)
        
        # Match citations to references
        uncited_refs, missing_refs = match_citations_to_references(intext_citations, reference_citations)
        
        results[code] = {
            'uncited_refs': uncited_refs,
            'missing_refs': missing_refs,
            'total_citations': len(intext_citations),
            'total_references': len(reference_citations),
            'has_issues': len(uncited_refs) > 0 or len(missing_refs) > 0
        }
    
    return results


# =============================================================================
# COMPONENT 3: Individual Reference Splitting
# =============================================================================

def split_references(references_text, citation_style='APA'):
    """
    Split references section into individual entries.
    Uses style-appropriate logic.
    """
    if citation_style == 'MLA':
        return split_mla_references(references_text)
    else:
        return split_apa_references(references_text)


def split_mla_references(references_text):
    """
    Split MLA references into individual entries.
    
    MLA Style Guide:
    - References start with "LastName, FirstName." or "LastName, First, et al."
    - Multiple authors: "LastName, FirstName, and FirstName LastName."
    - Can also start with "Title" if no author
    - Sections separated by periods
    - Year typically at end
    - Skip page headers: "LastName PageNumber"
    - Detects new references that start mid-line after URLs
    - Prevents splitting on middle initials (e.g., "Robert G.")
    """
    if not references_text:
        return []
    
    references = []
    current_ref = []
    
    for line in references_text:
        # Check for hanging indent BEFORE stripping
        # MLA uses hanging indents: first line is left-justified, continuation lines are indented
        is_indented = line.startswith(' ') or line.startswith('\t')
        
        line = line.strip()
        if not line:
            continue
        
        # Skip MLA page headers: "LastName PageNumber"
        if re.match(r'^[A-Z][a-z]+\s+\d+$', line):
            continue
        
        # If line is indented (hanging indent), it's a continuation of the current reference
        if is_indented and current_ref:
            current_ref.append(line)
            continue
        
        # Check if line contains a new reference starting mid-line
        # Pattern: URL/period followed by space and new author pattern
        # Example: "...scu.edu/vol12/iss1/7/. McCarthy, Joseph. ..."
        # Look for: end of URL (with domain OR file extension) OR just period + space + author pattern
        # BUT: Don't split on middle initials (single capital letter + period)
        # Note: [a-zA-Z] allows for names with internal capitals like McCarthy, McDonald, O'Brien
        # Recognizes both domain endings (.com/, .edu/) and file extensions (.htm, .html, .pdf, .asp, etc.)
        mid_line_candidates = list(re.finditer(r'(\.(?:com|org|edu|gov|net)/[^\s]*\.?|\.(?:htm|html|pdf|asp|aspx|php|jsp)\s*\.?|\.)\s+([A-Z][a-zA-Z]{2,},\s+[A-Z][a-zA-Z]{2,})', line))
        
        split_performed = False
        for match in mid_line_candidates:
            split_point = match.start(2)
            
            # Check if the period is part of a middle initial
            # Look back 2-4 characters to see if we have " X." pattern (middle initial)
            if split_point >= 3:
                before_split = line[max(0, split_point-4):split_point]
                # Pattern: space + single capital + period (middle initial)
                if re.search(r'\s[A-Z]\.\s*$', before_split):
                    # This is a middle initial, don't split here
                    continue
            
            # Only split if the "before" part is substantial (> 40 chars)
            # This avoids splitting cases like "Ball, Lucille. Love, Lucy." where Love, Lucy is a book title
            # But still splits "...long URL... Author, Name." where it's two merged references
            before = line[:split_point].strip()
            if len(before) <= 40:
                continue
            
            # Found a valid split point - split the line
            after = line[split_point:].strip()
            
            if before and current_ref:
                current_ref.append(before)
                # Save current reference
                ref_text = ' '.join(current_ref)
                if len(ref_text) > 20:
                    references.append(ref_text)
                # Start new reference with the part after split
                current_ref = [after]
            elif before:
                # First reference on the line - save even if short (legitimate author names can be < 20 chars)
                # E.g., "Ball, Lucille." is only 14 characters but is a valid complete author line
                references.append(before)
                current_ref = [after]
            else:
                # Just start with the after part
                current_ref = [after]
            
            split_performed = True
            break
        
        if split_performed:
            continue
        
        is_new_ref = False
        if current_ref:
            # Strategy: Check if PREVIOUS reference looks complete, then check if CURRENT line
            # could plausibly start a new reference. This is more flexible than hard-coded patterns.
            
            # First, check if previous reference looks complete
            last_line = current_ref[-1] if current_ref else ""
            prev_looks_complete = (
                re.search(r'\b[12]\d{3}\s*\.?\s*$', last_line) or  # ends with year
                re.search(r'pp?\.\s*\d+[\d\s–-]*\.?\s*$', last_line) or  # ends with page numbers
                re.search(r'(?:com|org|edu|gov|net)/[^\s]*\.?\s*$', last_line) or  # ends with URL
                re.search(r'\.(?:htm|html|pdf|asp|aspx|php|jsp)\.?\s*$', last_line) or  # ends with file extension
                re.search(r'Accessed\s+\d+\s+\w+\.?\s+\d{4}\.?\s*$', last_line) or  # ends with access date
                re.search(r'vol\.\s*\d+.*\d+\.?\s*$', last_line) or  # ends with volume info
                re.search(r'\d+\s+ed\.?\s*$', last_line) or  # ends with edition
                re.search(r'[a-zA-Z0-9\-/]{3,}\.\s*$', last_line)  # ends with word/URL segment + period (allows hyphens, numbers, slashes for URLs)
            )
            
            # Now check if current line could plausibly start a reference
            # Pattern 1: Starts with opening quote (title-first, no author)
            if re.match(r'^["\u201c]', line):
                is_new_ref = True
            
            # Pattern 2: Author patterns (has comma in the right place)
            # "LastName, FirstName" or "LastName, FirstName, et al." or "LastName, FirstName, and..."
            # Allow hyphens and apostrophes in names: O'Brien, McDonald-Smith, Kirshenblatt-Gimblett
            elif re.match(r'^[A-Z][a-zA-Z\-\']+,\s+[A-Z]', line):
                # Special case: "United States, Congress" needs completeness check
                if re.match(r'^United States,', line):
                    is_new_ref = prev_looks_complete
                # Special case: "States, Congress" is part of "United States"
                elif re.match(r'^States,', line):
                    is_new_ref = False
                # All other "LastName, FirstName" patterns
                else:
                    is_new_ref = prev_looks_complete
            
            # Pattern 3: Title-first references (multi-word titles OR single-word titles with period)
            # Multi-word: "The Red Menace", "A Collection of"
            # Single-word: "Requiem." or "Starry" (followed by more text)
            # This catches films, artworks, books without authors, etc.
            elif re.match(r'^[A-Z][a-zA-Z]*\s+[A-Z]', line) or re.match(r'^[A-Z][a-zA-Z]+\.\s+[A-Z]', line):
                is_new_ref = prev_looks_complete
            
            # If none of the above patterns match, it's a continuation
        
        if is_new_ref and current_ref:
            ref_text = ' '.join(current_ref)
            if len(ref_text) > 20:
                references.append(ref_text)
            current_ref = [line]
        else:
            current_ref.append(line)
    
    if current_ref:
        ref_text = ' '.join(current_ref)
        if len(ref_text) > 20:
            references.append(ref_text)
    
    return references


def split_apa_references(references_text):
    """
    Split APA references into individual entries.
    
    APA Style Guide:
    - References start with "Author, I. N. (Year)."
    - Year in parentheses early in reference
    - Initials, not full first names
    - Handles cases where URL ends with / and next reference starts on same line
    """
    if not references_text:
        return []
    
    references = []
    current_ref = []
    
    for line in references_text:
        line = line.strip()
        if not line:
            continue
        
        # Check if line contains a new reference starting mid-line
        # Pattern: URL ending with / followed by space and APA author pattern
        # Example: "...PMC7450866/ Author, I. N. (2020)..."
        # APA author pattern: "LastName, I." or "LastName, I. N."
        mid_line_candidates = list(re.finditer(r'(/)\s+([A-Z][a-zA-Z\-\']+,\s+[A-Z]\.(?:\s+[A-Z]\.)?)', line))
        
        split_performed = False
        for match in mid_line_candidates:
            split_point = match.start(2)
            
            # Only split if the "before" part is substantial (> 40 chars)
            # This avoids splitting legitimate references that happen to have "/ Author"
            before = line[:split_point].strip()
            if len(before) <= 40:
                continue
            
            # Check if the after part looks like an APA reference (has year in parentheses)
            after = line[split_point:].strip()
            if not re.search(r'\(\d{4}[a-z]?\)', after[:100]):
                continue
            
            # Found a valid split point - split the line
            if before and current_ref:
                current_ref.append(before)
                # Save current reference
                ref_text = ' '.join(current_ref)
                if len(ref_text) > 20:
                    references.append(ref_text)
                # Start new reference with the part after split
                current_ref = [after]
            elif before:
                # First reference on the line
                if len(before) > 20:
                    references.append(before)
                current_ref = [after]
            else:
                # Just start with the after part
                current_ref = [after]
            
            split_performed = True
            break
        
        if split_performed:
            continue
        
        is_new_ref = False
        if current_ref:
            if re.match(r'^[A-Z]', line):
                first_part = line[:150]
                
                # APA patterns for new reference:
                # 1. Year in parentheses: "(2020)"
                if re.search(r'\(\d{4}[a-z]?\)', first_part):
                    is_new_ref = True
                # 2. Year with comma/period early: "2020." or "2020,"
                elif re.search(r'\b(19|20)\d{2}\b', first_part) and re.search(r'[,\.]', first_part[:50]):
                    is_new_ref = True
        
        if is_new_ref and current_ref:
            ref_text = ' '.join(current_ref)
            if len(ref_text) > 20:
                references.append(ref_text)
            current_ref = [line]
        else:
            current_ref.append(line)
    
    if current_ref:
        ref_text = ' '.join(current_ref)
        if len(ref_text) > 20:
            references.append(ref_text)
    
    return references


def split_references_from_results(results):
    """Split references for all papers using style-appropriate logic."""
    output = {}
    for code, info in results.items():
        citation_style = info.get('citation_style', 'APA')
        individual_refs = split_references(info['references_text'], citation_style=citation_style)
        numbered_refs = [{'ref_num': i, 'text': ref} for i, ref in enumerate(individual_refs, 1)]
        output[code] = {
            'student_name': info['student_name'],
            'filepath': info['filepath'],
            'references': numbered_refs,
            'citation_style': citation_style,
            'style_confidence': info.get('style_confidence', 0)
        }
    return output


# =============================================================================
# COMPONENT 4: Reference Parsing
# =============================================================================

def parse_reference(ref_text, citation_style='APA'):
    """
    Parse a reference into components.
    Handles both APA and MLA styles.
    
    APA: Author, I. N. (Year). Title. Source.
    MLA: Author, Full Name. "Title." Container, elements, year.
    """
    if citation_style == 'MLA':
        return parse_mla_reference(ref_text)
    else:
        return parse_apa_reference(ref_text)


def parse_mla_reference(ref_text):
    """
    Parse an MLA reference into components.
    
    MLA Structure:
    1. Author. (or starts with "Title" if no author)
    2. "Title of Source."
    3. Container with comma-separated elements, ending with year.
    """
    result = {
        'raw': ref_text,
        'authors': None,
        'year': None,
        'title': None,
        'source': None,
        'doi': None,
        'ref_type': 'other'
    }
    
    text_lower = ref_text.lower()
    
    # Detect reference type (MLA-specific patterns)
    # Check journal indicators FIRST (before generic http check)
    if re.search(r'vol\.\s*\d+', text_lower) or re.search(r',\s*no\.\s*\d+', text_lower) or re.search(r'pp\.\s*\d+', text_lower):
        result['ref_type'] = 'journal'
    # DOI links are journals, not websites
    elif 'doi.org' in text_lower:
        result['ref_type'] = 'journal'
    # Now check for generic websites
    elif 'www.' in text_lower or 'http' in text_lower:
        result['ref_type'] = 'website'
    # Check for book indicators
    elif any(pub in text_lower for pub in ['press', 'publisher', 'publishing', 'books']):
        result['ref_type'] = 'book'
    
    # Extract DOI if present
    doi_match = re.search(r'(10\.\d{4,}/[^\s]+)', ref_text)
    if doi_match:
        result['doi'] = doi_match.group(1).rstrip('.,')
    
    # MLA structure: Split by periods to find major sections
    # Be careful with periods in titles (in quotes) and abbreviations
    
    # 1. Extract author (everything before first period, if it looks like a name)
    # Pattern: "LastName, FirstName." or starts with quote (no author)
    if ref_text.startswith('"') or ref_text.startswith('"'):
        # No author - starts with title
        result['authors'] = ''
    else:
        # Look for author pattern: "LastName, FirstName" followed by period
        author_match = re.match(r'^([^"]+?)\.\s+["\"]', ref_text)
        if author_match:
            result['authors'] = author_match.group(1).strip()
        else:
            # Fallback: take everything before first period
            first_period = ref_text.find('.')
            if first_period > 0 and first_period < 100:  # Reasonable author length
                result['authors'] = ref_text[:first_period].strip()
    
    # 2. Extract title (in quotes)
    title_match = re.search(r'["\u201c]([^"\u201d]+)["\u201d]', ref_text)
    if title_match:
        result['title'] = title_match.group(1).strip()
    
    # 3. Extract year (usually at the end, 4 digits)
    # Look for year in last 50 characters
    year_match = re.search(r'\b(19|20)\d{2}\b', ref_text[-50:])
    if year_match:
        result['year'] = year_match.group(0)
    
    # 4. Extract source/container (after title, before year)
    if title_match and result['year']:
        # Everything between title and year
        after_title = ref_text[title_match.end():].strip()
        year_pos = after_title.rfind(result['year'])
        if year_pos > 0:
            result['source'] = after_title[:year_pos].strip().rstrip(',.')
    elif title_match:
        # Just take what's after title
        result['source'] = ref_text[title_match.end():].strip().rstrip(',.')
    
    return result


def parse_apa_reference(ref_text):
    """Parse an APA reference into components."""
    result = {
        'raw': ref_text,
        'authors': None,
        'year': None,
        'title': None,
        'source': None,
        'doi': None,
        'ref_type': 'other'
    }
    
    text_lower = ref_text.lower()
    
    # Detect reference type
    if re.search(r'\bIn\s+[A-Z].*\(Ed[s]?\.\)', ref_text):
        result['ref_type'] = 'chapter'
    elif 'retrieved from' in text_lower or ('http' in text_lower and 'doi.org' not in text_lower):
        result['ref_type'] = 'website'
    elif re.search(r'\d+\(\d+\)', ref_text) or re.search(r'vol\.\s*\d+', text_lower):
        result['ref_type'] = 'journal'
    elif re.search(r'\.\s*[A-Z][A-Za-z\s&\-]+\s*(Press|Publishers?|Publications?|Books?|Publishing|Inc\.?|LLC|Company|Co\.?)\.?\s*$', ref_text):
        result['ref_type'] = 'book'
    elif re.search(r'\.\s*(Jossey-Bass|Wiley|Springer|Elsevier|Sage|Routledge|McGraw-Hill|Pearson|Cambridge|Oxford|Harvard|MIT|Yale|Stanford|Norton|Penguin|Random House|Simon & Schuster|HarperCollins|Macmillan|Houghton Mifflin|Cengage|Guilford|Erlbaum|Psychology Press|Academic Press|Shambhala|New Harbinger|Bantam|Vintage|Knopf)\.?\s*$', ref_text, re.IGNORECASE):
        result['ref_type'] = 'book'
    elif re.search(r'[A-Z][a-z]+,\s*[A-Z]{2}:', ref_text) or re.search(r':\s*[A-Z][a-z]+\s+(Press|Publishers?|Publications?|Books?|Publishing)', ref_text):
        result['ref_type'] = 'book'
    elif re.search(r'\(\d+(st|nd|rd|th)\s+ed\.\)', text_lower):
        result['ref_type'] = 'book'
    elif not re.search(r'\d+\(\d+\)', ref_text) and 'http' not in text_lower and re.search(r'\.\s*[A-Z][A-Za-z\-]+\.?\s*$', ref_text):
        result['ref_type'] = 'book'
    
    # Extract DOI
    doi_match = re.search(r'(10\.\d{4,}/[^\s]+)', ref_text)
    if doi_match:
        result['doi'] = doi_match.group(1).rstrip('.,')
    
    # Extract year
    year_match = re.search(r'\((\d{4})[a-z]?\)', ref_text)
    if year_match:
        result['year'] = year_match.group(1)
    else:
        year_match = re.search(r'\b(19|20)(\d{2})\b', ref_text)
        if year_match:
            result['year'] = year_match.group(0)
    
    # Extract authors
    if result['year']:
        year_pos = ref_text.find(result['year'])
        if year_pos > 0:
            authors = ref_text[:year_pos].strip()
            authors = re.sub(r'[\(\.,\s]+$', '', authors)
            if authors:
                result['authors'] = authors
    
    # Extract title
    year_pattern = re.search(r'\(\d{4}[a-z]?\)\.\s*', ref_text)
    if year_pattern:
        after_year = ref_text[year_pattern.end():]
        title_match = re.match(r'^([^\.]+\.)', after_year)
        if title_match:
            result['title'] = title_match.group(1).strip()
    
    if not result['title']:
        quote_match = re.search(r'["\u201c](.+)["\u201d][,\.]', ref_text)
        if quote_match:
            result['title'] = quote_match.group(1)
        else:
            quote_match = re.search(r'["\u201c]([^"\u201c]{10,})["\u201d]', ref_text)
            if quote_match:
                result['title'] = quote_match.group(1)
    
    return result


def parse_all_references(split_results):
    """Parse all references from split results."""
    output = {}
    for code, info in split_results.items():
        citation_style = info.get('citation_style', 'APA')
        parsed_refs = []
        for ref in info['references']:
            parsed = parse_reference(ref['text'], citation_style=citation_style)
            parsed['ref_num'] = ref['ref_num']
            parsed_refs.append(parsed)
        output[code] = {
            'student_name': info['student_name'],
            'filepath': info['filepath'],
            'references': parsed_refs,
            'citation_style': citation_style,
            'style_confidence': info.get('style_confidence', 0)
        }
    return output


# =============================================================================
# COMPONENT 5: Verification
# =============================================================================

def verify_by_doi(doi):
    """Look up a DOI via CrossRef."""
    try:
        response = requests.get(f"{CROSSREF_API}/{doi}", headers=HEADERS, timeout=10)
        if response.status_code == 200:
            work = response.json().get('message', {})
            title = work.get('title', [''])[0] if work.get('title') else ''
            authors = [f"{a.get('family', '')}, {a.get('given', '')}" for a in work.get('author', [])]
            year = None
            for key in ['published-print', 'published-online', 'issued']:
                if work.get(key):
                    year = work[key].get('date-parts', [[None]])[0][0]
                    if year:
                        break
            return {
                'found': True,
                'metadata': {
                    'title': title,
                    'authors': authors,
                    'year': str(year) if year else None,
                    'journal': work.get('container-title', [''])[0] if work.get('container-title') else ''
                },
                'error': None
            }
        elif response.status_code == 404:
            return {'found': False, 'metadata': None, 'error': 'DOI not found'}
        else:
            return {'found': False, 'metadata': None, 'error': f'HTTP {response.status_code}'}
    except Exception as e:
        return {'found': False, 'metadata': None, 'error': str(e)}


def search_by_title(title, author=None):
    """Search CrossRef by title."""
    try:
        clean_title = re.sub(r'[^\w\s]', ' ', title)
        clean_title = ' '.join(clean_title.split())
        params = {'query.title': clean_title, 'rows': 5}
        if author:
            params['query.author'] = author.split(',')[0].strip()
        
        response = requests.get(CROSSREF_API, params=params, headers=HEADERS, timeout=10)
        if response.status_code == 200:
            items = response.json().get('message', {}).get('items', [])
            matches = []
            for item in items:
                item_title = item.get('title', [''])[0] if item.get('title') else ''
                year = None
                for key in ['published-print', 'published-online', 'issued']:
                    if item.get(key):
                        year = item[key].get('date-parts', [[None]])[0][0]
                        if year:
                            break
                matches.append({
                    'title': item_title,
                    'year': str(year) if year else None,
                    'doi': item.get('DOI'),
                    'source': 'crossref'
                })
            return {'found': len(matches) > 0, 'matches': matches, 'error': None}
        return {'found': False, 'matches': [], 'error': f'HTTP {response.status_code}'}
    except Exception as e:
        return {'found': False, 'matches': [], 'error': str(e)}


def search_open_library(title, author=None):
    """Search Open Library for books."""
    try:
        clean_title = re.sub(r'[^\w\s]', ' ', title)
        query = f'title:{" ".join(clean_title.split())}'
        if author:
            query += f' author:{author.split(",")[0].strip()}'
        
        response = requests.get(
            'https://openlibrary.org/search.json',
            params={'q': query, 'limit': 5},
            headers=HEADERS,
            timeout=10
        )
        if response.status_code == 200:
            docs = response.json().get('docs', [])
            matches = [{
                'title': doc.get('title', ''),
                'year': str(doc.get('first_publish_year')) if doc.get('first_publish_year') else None,
                'doi': None,
                'source': 'openlibrary'
            } for doc in docs]
            return {'found': len(matches) > 0, 'matches': matches, 'error': None}
        return {'found': False, 'matches': [], 'error': f'HTTP {response.status_code}'}
    except Exception as e:
        return {'found': False, 'matches': [], 'error': str(e)}


def compare_titles(title1, title2):
    """Compare two titles for similarity (0-1)."""
    if not title1 or not title2:
        return 0
    
    def normalize(t):
        t = t.lower()
        t = re.sub(r'["\'\u201c\u201d\u2018\u2019"\'«»]', '', t)
        t = re.sub(r'[^\w\s]', '', t)
        return t.split()
    
    t1, t2 = normalize(title1), normalize(title2)
    if not t1 or not t2:
        return 0
    
    set1, set2 = set(t1), set(t2)
    intersection = len(set1 & set2)
    union = len(set1 | set2)
    return intersection / union if union > 0 else 0


def search_pubmed(title, author=None):
    """Search PubMed for journal articles."""
    try:
        # Clean title for search
        clean_title = re.sub(r'[^\w\s]', ' ', title)
        clean_title = ' '.join(clean_title.split())
        
        # Build search query
        query = f'{clean_title}[Title]'
        if author:
            author_last = author.split(',')[0].strip()
            query += f' AND {author_last}[Author]'
        
        # First, search for matching articles
        search_url = 'https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi'
        search_params = {
            'db': 'pubmed',
            'term': query,
            'retmax': 5,
            'retmode': 'json'
        }
        
        response = requests.get(search_url, params=search_params, headers=HEADERS, timeout=10)
        if response.status_code != 200:
            return {'found': False, 'matches': [], 'error': f'HTTP {response.status_code}'}
        
        result = response.json()
        id_list = result.get('esearchresult', {}).get('idlist', [])
        
        if not id_list:
            return {'found': False, 'matches': [], 'error': None}
        
        # Fetch details for found articles
        summary_url = 'https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi'
        summary_params = {
            'db': 'pubmed',
            'id': ','.join(id_list),
            'retmode': 'json'
        }
        
        response = requests.get(summary_url, params=summary_params, headers=HEADERS, timeout=10)
        if response.status_code != 200:
            return {'found': False, 'matches': [], 'error': f'HTTP {response.status_code}'}
        
        summary_result = response.json().get('result', {})
        
        matches = []
        for pmid in id_list:
            if pmid in summary_result:
                article = summary_result[pmid]
                pub_date = article.get('pubdate', '')
                year = pub_date[:4] if pub_date else None
                
                matches.append({
                    'title': article.get('title', ''),
                    'year': year,
                    'authors': [a.get('name', '') for a in article.get('authors', [])],
                    'doi': article.get('elocationid', '').replace('doi: ', '') if 'doi' in article.get('elocationid', '').lower() else None,
                    'source': 'pubmed',
                    'pmid': pmid
                })
        
        return {'found': len(matches) > 0, 'matches': matches, 'error': None}
    
    except Exception as e:
        return {'found': False, 'matches': [], 'error': str(e)}


def search_google_books(title, author=None):
    """Search Google Books API for books."""
    try:
        clean_title = re.sub(r'[^\w\s]', ' ', title)
        query = f'intitle:{" ".join(clean_title.split())}'
        if author:
            query += f'+inauthor:{author.split(",")[0].strip()}'
        
        response = requests.get(
            'https://www.googleapis.com/books/v1/volumes',
            params={'q': query, 'maxResults': 5},
            headers=HEADERS,
            timeout=10
        )
        if response.status_code == 200:
            items = response.json().get('items', [])
            matches = []
            for item in items:
                vol_info = item.get('volumeInfo', {})
                matches.append({
                    'title': vol_info.get('title', ''),
                    'year': vol_info.get('publishedDate', '')[:4] if vol_info.get('publishedDate') else None,
                    'authors': vol_info.get('authors', []),
                    'doi': None,
                    'source': 'google_books',
                    'info_link': vol_info.get('infoLink', '')
                })
            return {'found': len(matches) > 0, 'matches': matches, 'error': None}
        return {'found': False, 'matches': [], 'error': f'HTTP {response.status_code}'}
    except Exception as e:
        return {'found': False, 'matches': [], 'error': str(e)}


def verify_reference(ref, verified_threshold=0.95, partial_threshold=0.70, ignore_websites=False):
    """Verify a single reference."""
    ref_type = ref.get('ref_type', 'other')
    
    if ref_type == 'website':
        if ignore_websites:
            return {'status': 'website_skipped', 'message': 'Website skipped per settings', 'crossref_data': None}
        return {'status': 'website_manual_verify', 'message': 'Website detected - verify manually', 'crossref_data': None}
    
    if ref_type in ('book', 'chapter'):
        if ref.get('title'):
            # Try Open Library first
            ol_result = search_open_library(ref['title'], ref.get('authors'))
            if ol_result['found'] and ol_result['matches']:
                best = ol_result['matches'][0]
                sim = compare_titles(best['title'], ref['title'])
                if sim > verified_threshold:
                    return {'status': 'verified', 'message': f'Book found in Open Library ({sim:.0%})', 'crossref_data': best}
                elif sim > partial_threshold:
                    return {'status': 'partial_match', 'message': f'Partial book match ({sim:.0%})', 'crossref_data': best, 'student_title': ref['title'], 'crossref_title': best['title']}
            
            # Try Google Books as backup
            gb_result = search_google_books(ref['title'], ref.get('authors'))
            if gb_result['found'] and gb_result['matches']:
                best = gb_result['matches'][0]
                sim = compare_titles(best['title'], ref['title'])
                if sim > verified_threshold:
                    return {'status': 'verified', 'message': f'Book found in Google Books ({sim:.0%})', 'crossref_data': best}
                elif sim > partial_threshold:
                    return {'status': 'partial_match', 'message': f'Partial book match ({sim:.0%})', 'crossref_data': best, 'student_title': ref['title'], 'crossref_title': best['title']}
        
        # Generate search URL for manual verification using full reference
        search_text = ref.get('raw', ref.get('title', ''))[:150]  # Use raw reference, limit length
        search_query = re.sub(r'[^\w\s]', ' ', search_text)
        search_query = ' '.join(search_query.split())  # Normalize whitespace
        search_url = f"https://www.google.com/search?q={search_query.replace(' ', '+')}"
        return {'status': 'book_manual_verify', 'message': 'Book/chapter not found - verify manually', 'crossref_data': None, 'search_url': search_url}
    
    if ref.get('doi'):
        result = verify_by_doi(ref['doi'])
        if result['found']:
            cr_title = result['metadata']['title']
            ref_title = ref.get('title', '')
            raw_text = ref.get('raw', '')
            sim = compare_titles(cr_title, ref_title)
            raw_sim = compare_titles(cr_title, raw_text)
            if sim > 0.3 or raw_sim > 0.5:
                return {'status': 'verified', 'message': f'DOI verified ({max(sim, raw_sim):.0%})', 'crossref_data': result['metadata']}
            else:
                return {'status': 'doi_mismatch', 'message': f'DOI exists but title doesn\'t match. CrossRef: "{cr_title[:80]}"', 'crossref_data': result['metadata'], 'student_title': ref_title or raw_text[:100], 'crossref_title': cr_title}
        return {'status': 'no_match', 'message': f'DOI not found: {ref["doi"]}', 'crossref_data': None}
    
    if ref.get('title'):
        # Try CrossRef first
        result = search_by_title(ref['title'], ref.get('authors'))
        if result['found'] and result['matches']:
            best = result['matches'][0]
            sim = compare_titles(best['title'], ref['title'])
            if sim > verified_threshold:
                return {'status': 'verified', 'message': f'Title match found in CrossRef ({sim:.0%})', 'crossref_data': best}
            elif sim > partial_threshold:
                return {'status': 'partial_match', 'message': f'Partial match in CrossRef ({sim:.0%})', 'crossref_data': best, 'student_title': ref['title'], 'crossref_title': best['title']}
        
        # Try PubMed as fallback
        pubmed_result = search_pubmed(ref['title'], ref.get('authors'))
        if pubmed_result['found'] and pubmed_result['matches']:
            best = pubmed_result['matches'][0]
            sim = compare_titles(best['title'], ref['title'])
            if sim > verified_threshold:
                return {'status': 'verified', 'message': f'Title match found in PubMed ({sim:.0%})', 'crossref_data': best}
            elif sim > partial_threshold:
                return {'status': 'partial_match', 'message': f'Partial match in PubMed ({sim:.0%})', 'crossref_data': best, 'student_title': ref['title'], 'crossref_title': best['title']}
        
        return {'status': 'no_match', 'message': 'No matching publication found in CrossRef or PubMed', 'crossref_data': None}
    
    return {'status': 'partial_match', 'message': 'Could not parse enough information', 'crossref_data': None}


def verify_all_references(parsed_results, delay=0.3, verified_threshold=0.95, partial_threshold=0.70, ignore_books=False, ignore_websites=False):
    """Verify all references."""
    output = {}
    for code, info in parsed_results.items():
        verified_refs = []
        for ref in info['references']:
            # Skip books if ignore_books is enabled
            if ignore_books and ref.get('ref_type') in ('book', 'chapter'):
                ref['verification'] = {
                    'status': 'skipped',
                    'message': 'Book/chapter skipped (ignore books enabled)',
                    'crossref_data': None
                }
            else:
                verification = verify_reference(ref, verified_threshold, partial_threshold, ignore_websites)
                ref['verification'] = verification
                time.sleep(delay)
            verified_refs.append(ref)
        output[code] = {
            'student_name': info['student_name'],
            'filepath': info['filepath'],
            'references': verified_refs,
            'citation_style': info.get('citation_style', 'UNKNOWN'),
            'style_confidence': info.get('style_confidence', 0)
        }
    return output


# =============================================================================
# COMPONENT 6: Report Generation
# =============================================================================

def generate_report(verified_results, output_path, verified_threshold=95, partial_threshold=70, ignore_books=False, ignore_websites=False, citation_results=None):
    """Generate a DOCX report."""
    doc = DocxDocument()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    title = doc.add_heading('Reference Verification Report', 0)
    title.runs[0].font.color.rgb = FT_CRIMSON_RGB
    
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.italic = True
    
    # Calculate stats
    stats = {'total_papers': len(verified_results), 'total_refs': 0, 'verified': 0, 'no_match': 0, 
             'partial_match': 0, 'doi_mismatch': 0, 'book_manual': 0, 'website_manual': 0, 'skipped': 0,
             'citation_total_citations': 0, 'citation_total_references': 0, 
             'citation_uncited': 0, 'citation_missing': 0}
    
    for code, info in verified_results.items():
        for ref in info['references']:
            stats['total_refs'] += 1
            status = ref['verification']['status']
            if status == 'verified': stats['verified'] += 1
            elif status == 'no_match': stats['no_match'] += 1
            elif status == 'partial_match': stats['partial_match'] += 1
            elif status == 'doi_mismatch': stats['doi_mismatch'] += 1
            elif status == 'book_manual_verify': stats['book_manual'] += 1
            elif status == 'website_manual_verify': stats['website_manual'] += 1
            elif status == 'website_skipped': stats['skipped'] += 1
            elif status == 'skipped': stats['skipped'] += 1
    
    # Calculate citation matching stats if available
    if citation_results:
        for code, cite_info in citation_results.items():
            stats['citation_total_citations'] += cite_info['total_citations']
            stats['citation_total_references'] += cite_info['total_references']
            stats['citation_uncited'] += len(cite_info['uncited_refs'])
            stats['citation_missing'] += len(cite_info['missing_refs'])
    
    # Summary
    doc.add_heading('Summary', level=1)
    
    summary_data = [
        ('Total Papers', str(stats['total_papers']), None),
        ('Total References', str(stats['total_refs']), None),
        ('Verified', str(stats['verified']), None),
        ('No Match (Suspicious)', str(stats['no_match']), FT_CRIMSON_RGB if stats['no_match'] > 0 else None),
        ('Partial Match', str(stats['partial_match']), FT_ORANGE_RGB if stats['partial_match'] > 0 else None),
        ('DOI Mismatch', str(stats['doi_mismatch']), FT_CRIMSON_RGB if stats['doi_mismatch'] > 0 else None),
    ]
    
    if ignore_books and ignore_websites:
        summary_data.append(('Books & Websites (Skipped)', str(stats['skipped']), None))
    elif ignore_books:
        summary_data.append(('Books (Skipped)', str(stats['skipped']), None))
        summary_data.append(('Websites (Manual)', str(stats['website_manual']), None))
    elif ignore_websites:
        summary_data.append(('Books (Manual)', str(stats['book_manual']), None))
        summary_data.append(('Websites (Skipped)', str(stats['skipped']), None))
    else:
        summary_data.append(('Book/Chapter (Manual)', str(stats['book_manual']), None))
        summary_data.append(('Website (Manual)', str(stats['website_manual']), None))
    
    # Add citation matching stats if available
    if citation_results:
        summary_data.extend([
            ('', '', None),  # Spacer
            ('--- Citation-Reference Matching ---', '', None),
            ('Total In-Text Citations', str(stats['citation_total_citations']), None),
            ('Total Reference Entries', str(stats['citation_total_references']), None),
            ('Uncited References', str(stats['citation_uncited']), FT_ORANGE_RGB if stats['citation_uncited'] > 0 else None),
            ('Missing References', str(stats['citation_missing']), FT_CRIMSON_RGB if stats['citation_missing'] > 0 else None),
        ])
    
    summary_table = doc.add_table(rows=len(summary_data), cols=2)
    summary_table.style = 'Table Grid'
    
    for i, (label, value, color) in enumerate(summary_data):
        row = summary_table.rows[i]
        if label == '':
            # Skip empty rows
            row.cells[0].text = ''
            row.cells[1].text = ''
        elif label.startswith('---'):
            # Section header
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = FT_CRIMSON_RGB
            row.cells[1].text = ''
        else:
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True
            if color:
                row.cells[1].paragraphs[0].runs[0].font.color.rgb = color
                row.cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Legend
    doc.add_heading('Status Definitions', level=2)
    legend_items = [
        ('Verified', f'Reference found with ≥{verified_threshold}% title match.'),
        ('No Match', 'Reference not found - may be fabricated or incorrectly cited.'),
        ('Partial Match', f'Similar reference found ({partial_threshold}-{verified_threshold}% match).'),
        ('DOI Mismatch', 'DOI link resolves to a different paper than cited.'),
    ]
    if ignore_books and ignore_websites:
        legend_items.append(('Skipped', 'Book/chapter/website skipped per settings.'))
    elif ignore_books:
        legend_items.append(('Skipped', 'Book/chapter skipped per settings.'))
        legend_items.append(('Website (Manual)', 'Website source - requires manual verification.'))
    elif ignore_websites:
        legend_items.append(('Book/Chapter (Manual)', 'Book or chapter - requires manual verification.'))
        legend_items.append(('Skipped', 'Website skipped per settings.'))
    else:
        legend_items.append(('Book/Chapter (Manual)', 'Book or chapter - requires manual verification.'))
        legend_items.append(('Website (Manual)', 'Website source - requires manual verification.'))
    
    for term, definition in legend_items:
        para = doc.add_paragraph()
        para.add_run(f'{term}: ').bold = True
        para.add_run(definition)
    
    doc.add_paragraph()
    
    # Details
    doc.add_heading('Detailed Findings by Student', level=1)
    for code, info in verified_results.items():
        student_heading = doc.add_heading(f'{info["student_name"]}', level=2)
        student_heading.runs[0].font.color.rgb = FT_CRIMSON_RGB
        
        # Add citation style detection info
        citation_style = info.get('citation_style', 'UNKNOWN')
        style_confidence = info.get('style_confidence', 0)
        
        style_para = doc.add_paragraph()
        style_para.add_run('Detected Style: ').bold = True
        style_para.add_run(f'{citation_style}  ')
        style_para.add_run('Confidence: ').bold = True
        confidence_run = style_para.add_run(f'{style_confidence:.0f}%')
        
        # Color code confidence
        if style_confidence >= 90:
            confidence_run.font.color.rgb = FT_GREEN_RGB
        elif style_confidence >= 70:
            confidence_run.font.color.rgb = FT_ORANGE_RGB
        else:
            confidence_run.font.color.rgb = FT_CRIMSON_RGB
        
        if style_confidence < 70 and citation_style != 'UNKNOWN':
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run('⚠ Low confidence - results may be inaccurate. Consider manually verifying style.')
            warning_run.font.color.rgb = FT_CRIMSON_RGB
            warning_run.font.size = Pt(10)
        
        refs = info['references']
        statuses = [r['verification']['status'] for r in refs]
        
        # Build summary line based on settings
        summary_parts = [
            f"Total: {len(refs)}",
            f"Verified: {statuses.count('verified')}",
            f"No Match: {statuses.count('no_match')}",
            f"Partial: {statuses.count('partial_match')}",
            f"DOI Mismatch: {statuses.count('doi_mismatch')}",
        ]
        if ignore_books and ignore_websites:
            summary_parts.append(f"Skipped: {statuses.count('skipped') + statuses.count('website_skipped')}")
        elif ignore_books:
            summary_parts.append(f"Skipped: {statuses.count('skipped')}")
            summary_parts.append(f"Website: {statuses.count('website_manual_verify')}")
        elif ignore_websites:
            summary_parts.append(f"Book: {statuses.count('book_manual_verify')}")
            summary_parts.append(f"Skipped: {statuses.count('website_skipped')}")
        else:
            summary_parts.append(f"Book: {statuses.count('book_manual_verify')}")
            summary_parts.append(f"Website: {statuses.count('website_manual_verify')}")
        
        summary_para = doc.add_paragraph(" | ".join(summary_parts))
        summary_para.runs[0].font.size = Pt(10)
        summary_para.runs[0].font.italic = True
        
        # Add citation matching results if available
        if citation_results and code in citation_results:
            cite_info = citation_results[code]
            doc.add_heading('Citation-Reference Matching', level=3)
            
            cite_summary_para = doc.add_paragraph()
            cite_summary_para.add_run(f"In-text citations found: {cite_info['total_citations']} | ")
            cite_summary_para.add_run(f"References in list: {cite_info['total_references']}")
            cite_summary_para.runs[0].font.size = Pt(10)
            
            # Report uncited references
            if cite_info['uncited_refs']:
                doc.add_paragraph()
                uncited_heading = doc.add_paragraph()
                uncited_run = uncited_heading.add_run(f"⚠ Uncited References ({len(cite_info['uncited_refs'])})")
                uncited_run.bold = True
                uncited_run.font.color.rgb = FT_ORANGE_RGB
                uncited_run.font.size = Pt(11)
                
                note_para = doc.add_paragraph("These references appear in the reference list but were not cited in the paper:")
                note_para.runs[0].font.size = Pt(9)
                note_para.runs[0].font.italic = True
                
                for authors, year in cite_info['uncited_refs']:
                    item_para = doc.add_paragraph(f"• {authors} ({year})", style='List Bullet')
                    item_para.paragraph_format.left_indent = Inches(0.25)
                    item_para.runs[0].font.size = Pt(10)
            
            # Report missing references
            if cite_info['missing_refs']:
                doc.add_paragraph()
                missing_heading = doc.add_paragraph()
                missing_run = missing_heading.add_run(f"⚠ Missing References ({len(cite_info['missing_refs'])})")
                missing_run.bold = True
                missing_run.font.color.rgb = FT_CRIMSON_RGB
                missing_run.font.size = Pt(11)
                
                note_para = doc.add_paragraph("These citations appear in the paper but have no matching reference:")
                note_para.runs[0].font.size = Pt(9)
                note_para.runs[0].font.italic = True
                
                for author, year in cite_info['missing_refs']:
                    item_para = doc.add_paragraph(f"• {author} ({year})", style='List Bullet')
                    item_para.paragraph_format.left_indent = Inches(0.25)
                    item_para.runs[0].font.size = Pt(10)
            
            # Success message if everything matches
            if not cite_info['uncited_refs'] and not cite_info['missing_refs']:
                doc.add_paragraph()
                success_para = doc.add_paragraph("✓ All in-text citations have matching references, and all references are cited.")
                success_para.runs[0].font.color.rgb = RGBColor(0, 102, 0)
                success_para.runs[0].font.size = Pt(10)
        
        # Filter out verified and skipped references
        problem_refs = [r for r in refs if r['verification']['status'] not in ('verified', 'skipped')]
        if problem_refs:
            doc.add_heading('References Needing Attention', level=3)
            for ref in problem_refs:
                v = ref['verification']
                status_para = doc.add_paragraph()
                status_run = status_para.add_run(f"[{v['status'].upper().replace('_', ' ')}] ")
                status_run.bold = True
                if v['status'] in ('no_match', 'doi_mismatch'):
                    status_run.font.color.rgb = FT_CRIMSON_RGB
                elif v['status'] == 'partial_match':
                    status_run.font.color.rgb = FT_ORANGE_RGB
                status_para.add_run(v['message']).font.size = Pt(10)
                
                ref_para = doc.add_paragraph()
                ref_para.paragraph_format.left_indent = Inches(0.25)
                ref_para.add_run(ref['raw']).font.size = Pt(10)
                
                # Add search URL for books to help manual verification
                if v['status'] == 'book_manual_verify':
                    # Use stored search_url or generate one from full reference
                    if v.get('search_url'):
                        search_url = v['search_url']
                    elif ref.get('raw'):
                        search_text = ref.get('raw', '')[:150]
                        search_query = re.sub(r'[^\w\s]', ' ', search_text)
                        search_query = ' '.join(search_query.split())
                        search_url = f"https://www.google.com/search?q={search_query.replace(' ', '+')}"
                    else:
                        search_url = None
                    
                    if search_url:
                        url_para = doc.add_paragraph()
                        url_para.paragraph_format.left_indent = Inches(0.25)
                        url_para.add_run('Search URL: ').bold = True
                        add_hyperlink(url_para, search_url, 'Click to search Google')
                
                if v.get('student_title') and v.get('crossref_title'):
                    compare_para = doc.add_paragraph()
                    compare_para.paragraph_format.left_indent = Inches(0.25)
                    compare_para.add_run('Student title: ').bold = True
                    compare_para.add_run(v['student_title']).font.size = Pt(9)
                    compare_para.add_run('\nDatabase title: ').bold = True
                    compare_para.add_run(v['crossref_title']).font.size = Pt(9)
                
                doc.add_paragraph()
        
        doc.add_page_break()
    
    doc.save(output_path)
    return stats


# =============================================================================
# GUI
# =============================================================================

class CrimsonButton(tk.Frame):
    """Custom button with proper colors on Mac."""
    def __init__(self, parent, text, command, **kwargs):
        super().__init__(parent, bg=FT_CRIMSON, padx=2, pady=2)
        self.label = tk.Label(self, text=text, bg=FT_CRIMSON, fg=FT_WHITE,
                              font=kwargs.get('font', ('Helvetica', 10)),
                              padx=kwargs.get('padx', 15), pady=kwargs.get('pady', 5), cursor='hand2')
        self.label.pack()
        self.command = command
        self.label.bind('<Button-1>', lambda e: self.on_click())
        self.label.bind('<Enter>', lambda e: self.label.config(bg='#990000'))
        self.label.bind('<Leave>', lambda e: self.label.config(bg=FT_CRIMSON))
    
    def on_click(self):
        if self.command:
            self.command()
    
    def config(self, **kwargs):
        if 'state' in kwargs:
            if kwargs['state'] == tk.DISABLED:
                self.label.config(bg='#CCCCCC', cursor='')
                self.label.unbind('<Button-1>')
            else:
                self.label.config(bg=FT_CRIMSON, cursor='hand2')
                self.label.bind('<Button-1>', lambda e: self.on_click())


class ReferenceCheckerGUI:
    def __init__(self, root):
        self.root = root
        self.root.title(f"Panther Reference Verification v{VERSION}")
        self.root.geometry("700x850")
        self.root.configure(bg=FT_WHITE)
        
        # Set window icon
        self.set_window_icon()
        
        # Store logo reference to prevent garbage collection
        self.logo_image = None
        
        self.input_folder = tk.StringVar()
        self.output_file = tk.StringVar(value="")
        self.status_text = tk.StringVar(value="Ready")
        self.verified_threshold = tk.StringVar(value="95")
        self.partial_threshold = tk.StringVar(value="70")
        self.ignore_books = tk.BooleanVar(value=False)
        self.ignore_websites = tk.BooleanVar(value=False)
        self.skip_citation_check = tk.BooleanVar(value=False)  # Default: do citation checking
        self.is_running = False
        
        # Track results widgets for clearing
        self.results_widgets = []
        
        self.create_menu()
        self.create_widgets()
        
        # Check for updates in background
        threading.Thread(target=self.check_updates_background, daemon=True).start()
    
    def set_window_icon(self):
        """Set the window icon based on platform."""
        try:
            icon_path = get_resource_path(os.path.join('assets', 'panther_icon.ico'))
            if os.path.exists(icon_path):
                if sys.platform == 'win32':
                    self.root.iconbitmap(icon_path)
                else:
                    # For Mac/Linux, use PhotoImage with PNG
                    png_path = get_resource_path(os.path.join('assets', 'panther_HQ.png'))
                    if os.path.exists(png_path):
                        icon = PhotoImage(file=png_path)
                        self.root.iconphoto(True, icon)
        except Exception:
            pass  # Silently fail if icon not found
    
    def create_menu(self):
        """Create the application menu bar."""
        menubar = Menu(self.root)
        self.root.config(menu=menubar)
        
        # Help menu
        help_menu = Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="User Guide", command=self.open_user_guide)
        help_menu.add_separator()
        help_menu.add_command(label="Check for Updates", command=self.check_updates_manual)
        help_menu.add_command(label="About", command=self.show_about)
    
    def open_user_guide(self):
        """Open the user guide PDF."""
        guide_path = get_resource_path(os.path.join('docs', 'Panther_Reference_Verification_User_Guide.pdf'))
        if os.path.exists(guide_path):
            webbrowser.open(f'file://{guide_path}')
        else:
            messagebox.showinfo("User Guide", 
                "The user guide is not yet available.\n\n"
                "For help, please visit the GitHub repository.")
    
    def show_about(self):
        """Show the About dialog."""
        messagebox.showinfo("About Panther Reference Verification",
            f"Panther Reference Verification\n"
            f"Version {VERSION}\n\n"
            f"Developed by Darby Proctor, Ph.D.\n\n"
            f"A tool to detect potentially fabricated\n"
            f"references in student papers.\n\n"
            f"Checks references against:\n"
            f"• CrossRef and PubMed (journals)\n"
            f"• Open Library and Google Books (books)")
    
    def check_updates_background(self):
        """Check for updates in background and show notification if available."""
        result = check_for_updates()
        if result and result[0]:  # has_update is True
            _, latest_version, download_url = result
            self.root.after(0, lambda: self.show_update_notification(latest_version, download_url))
    
    def check_updates_manual(self):
        """Manually check for updates."""
        result = check_for_updates()
        if result is None:
            messagebox.showerror("Update Check", "Could not check for updates.\nPlease check your internet connection.")
        elif result[0]:  # has_update
            _, latest_version, download_url = result
            self.show_update_notification(latest_version, download_url)
        else:
            messagebox.showinfo("Update Check", f"You are running the latest version (v{VERSION}).")
    
    def show_update_notification(self, latest_version, download_url):
        """Show update available notification."""
        response = messagebox.askyesno("Update Available",
            f"A new version is available!\n\n"
            f"Current version: v{VERSION}\n"
            f"Latest version: v{latest_version}\n\n"
            f"Would you like to download the update?")
        if response:
            webbrowser.open(download_url)
    
    def create_widgets(self):
        main_frame = tk.Frame(self.root, bg=FT_WHITE, padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Header frame with logo and title
        header_frame = tk.Frame(main_frame, bg=FT_WHITE)
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Logo on the left
        logo_frame = tk.Frame(header_frame, bg=FT_WHITE)
        logo_frame.pack(side=tk.LEFT, padx=(0, 15))
        
        try:
            # Try multiple possible locations for the logo
            possible_paths = [
                get_resource_path(os.path.join('assets', 'panther_HQ.png')),
                os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets', 'panther_HQ.png'),
                os.path.join(os.getcwd(), 'assets', 'panther_HQ.png'),
                'assets/panther_HQ.png',
                'panther_HQ.png',
            ]
            
            logo_path = None
            for path in possible_paths:
                if os.path.exists(path):
                    logo_path = path
                    break
            
            if logo_path:
                img = Image.open(logo_path)
                img = img.resize((60, 60), Image.Resampling.LANCZOS)
                self.logo_image = ImageTk.PhotoImage(img)
                logo_label = tk.Label(logo_frame, image=self.logo_image, bg=FT_WHITE)
                logo_label.pack()
        except Exception as e:
            print(f"Could not load logo: {e}")  # Debug info
        
        # Title and credit on the right of logo
        title_frame = tk.Frame(header_frame, bg=FT_WHITE)
        title_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        tk.Label(title_frame, text="Panther Reference Verification", font=('Helvetica', 24, 'bold'),
                 fg=FT_CRIMSON, bg=FT_WHITE, anchor='w').pack(anchor=tk.W)
        
        # Developer credit
        tk.Label(title_frame, text="Developed by Darby Proctor, Ph.D.",
                 font=('Helvetica', 9), fg='#666666', bg=FT_WHITE, anchor='w').pack(anchor=tk.W)
        
        # Important disclaimer box
        disclaimer_border = tk.Frame(main_frame, bg='#CC6600', padx=2, pady=2)
        disclaimer_border.pack(fill=tk.X, pady=(0, 20))
        disclaimer_frame = tk.Frame(disclaimer_border, bg='#FFF8DC', padx=15, pady=18)
        disclaimer_frame.pack(fill=tk.X)
        
        # Warning icon and title
        warning_header = tk.Frame(disclaimer_frame, bg='#FFF8DC')
        warning_header.pack(fill=tk.X)
        tk.Label(warning_header, text="⚠️", font=('Helvetica', 16), bg='#FFF8DC').pack(side=tk.LEFT, padx=(0, 5))
        tk.Label(warning_header, text="IMPORTANT: This is a Screening Tool", 
                 font=('Helvetica', 12, 'bold'), fg=FT_CRIMSON, bg='#FFF8DC').pack(side=tk.LEFT)
        
        # Disclaimer text
        disclaimer_text = (
            "This tool reduces your workload by automatically verifying references against academic databases.\n"
            "ALL flagged references still require manual verification by you (the instructor).\n\n"
            "Common legitimate reasons for flags: journal not indexed, honest student errors, recent publications"
        )
        tk.Label(disclaimer_frame, text=disclaimer_text, font=('Helvetica', 12), 
                 bg='#FFF8DC', fg='#333333', justify=tk.LEFT, wraplength=700).pack(anchor=tk.W, pady=(8, 8))
        
        # Input folder
        input_frame = tk.Frame(main_frame, bg=FT_WHITE)
        input_frame.pack(fill=tk.X, pady=10)
        tk.Label(input_frame, text="Papers Folder:", font=('Helvetica', 14), bg=FT_WHITE).pack(anchor=tk.W)
        input_row = tk.Frame(input_frame, bg=FT_WHITE)
        input_row.pack(fill=tk.X, pady=5)
        tk.Entry(input_row, textvariable=self.input_folder, font=('Helvetica', 13), width=50, state='readonly').pack(side=tk.LEFT, fill=tk.X, expand=True)
        CrimsonButton(input_row, text="Browse", command=self.browse_input, font=('Helvetica', 13)).pack(side=tk.LEFT, padx=(10, 0))
        
        # Output file
        output_frame = tk.Frame(main_frame, bg=FT_WHITE)
        output_frame.pack(fill=tk.X, pady=10)
        tk.Label(output_frame, text="Save Report As:", font=('Helvetica', 14), bg=FT_WHITE).pack(anchor=tk.W)
        output_row = tk.Frame(output_frame, bg=FT_WHITE)
        output_row.pack(fill=tk.X, pady=5)
        tk.Entry(output_row, textvariable=self.output_file, font=('Helvetica', 13), width=50, state='readonly').pack(side=tk.LEFT, fill=tk.X, expand=True)
        CrimsonButton(output_row, text="Browse", command=self.browse_output, font=('Helvetica', 13)).pack(side=tk.LEFT, padx=(10, 0))
        
        # Thresholds
        threshold_frame = tk.Frame(main_frame, bg=FT_WHITE)
        threshold_frame.pack(fill=tk.X, pady=10)
        tk.Label(threshold_frame, text="Matching Thresholds:", font=('Helvetica', 14), bg=FT_WHITE).pack(anchor=tk.W)
        threshold_row = tk.Frame(threshold_frame, bg=FT_WHITE)
        threshold_row.pack(fill=tk.X, pady=5)
        tk.Label(threshold_row, text="Verified ≥", font=('Helvetica', 13), bg=FT_WHITE).pack(side=tk.LEFT)
        tk.Entry(threshold_row, textvariable=self.verified_threshold, font=('Helvetica', 13), width=5).pack(side=tk.LEFT, padx=(5, 0))
        tk.Label(threshold_row, text="%", font=('Helvetica', 13), bg=FT_WHITE).pack(side=tk.LEFT, padx=(2, 20))
        tk.Label(threshold_row, text="Partial Match ≥", font=('Helvetica', 13), bg=FT_WHITE).pack(side=tk.LEFT)
        tk.Entry(threshold_row, textvariable=self.partial_threshold, font=('Helvetica', 13), width=5).pack(side=tk.LEFT, padx=(5, 0))
        tk.Label(threshold_row, text="%", font=('Helvetica', 13), bg=FT_WHITE).pack(side=tk.LEFT)
        
        # Options
        options_frame = tk.Frame(main_frame, bg=FT_WHITE)
        options_frame.pack(fill=tk.X, pady=5)
        tk.Checkbutton(options_frame, text="Ignore books/chapters (poor accuracy, many books must be manually verified)",
                       variable=self.ignore_books, font=('Helvetica', 13), bg=FT_WHITE,
                       activebackground=FT_WHITE).pack(anchor=tk.W)
        tk.Checkbutton(options_frame, text="Ignore websites (these must be manually verified)",
                       variable=self.ignore_websites, font=('Helvetica', 13), bg=FT_WHITE,
                       activebackground=FT_WHITE).pack(anchor=tk.W)
        tk.Checkbutton(options_frame, text="Skip citation-reference matching",
                       variable=self.skip_citation_check, font=('Helvetica', 13), bg=FT_WHITE,
                       activebackground=FT_WHITE).pack(anchor=tk.W)
        
        # Run button
        self.run_btn = CrimsonButton(main_frame, text="Run Verification", command=self.run_verification,
                                      font=('Helvetica', 16, 'bold'), padx=25, pady=10)
        self.run_btn.pack(pady=20)
        
        # Status
        status_border = tk.Frame(main_frame, bg=FT_CRIMSON, padx=2, pady=2)
        status_border.pack(fill=tk.X, pady=10)
        status_frame = tk.Frame(status_border, bg=FT_LIGHT_GRAY, padx=15, pady=15)
        status_frame.pack(fill=tk.X)
        tk.Label(status_frame, text="Status:", font=('Helvetica', 14, 'bold'), bg=FT_LIGHT_GRAY).pack(anchor=tk.W)
        tk.Label(status_frame, textvariable=self.status_text, font=('Helvetica', 14), bg=FT_LIGHT_GRAY,
                 wraplength=550, justify=tk.LEFT).pack(anchor=tk.W, pady=5)
        self.progress = ttk.Progressbar(status_frame, mode='indeterminate', length=500)
        self.progress.pack(pady=5)
        
        # Results frame
        self.results_frame = tk.Frame(main_frame, bg=FT_WHITE)
        self.results_frame.pack(fill=tk.BOTH, expand=True, pady=10)
    
    def browse_input(self):
        folder = filedialog.askdirectory(title="Select Folder with Student Papers")
        if folder:
            self.input_folder.set(folder)
    
    def browse_output(self):
        file = filedialog.asksaveasfilename(title="Save Report As", defaultextension=".docx",
                                            filetypes=[("Word Document", "*.docx")])
        if file:
            self.output_file.set(file)
    
    def update_status(self, message):
        self.status_text.set(message)
        self.root.update_idletasks()
    
    def clear_results(self):
        """Clear previous results display."""
        for widget in self.results_widgets:
            widget.destroy()
        self.results_widgets = []
    
    def run_verification(self):
        if self.is_running:
            return
        
        input_folder = self.input_folder.get()
        output_file = self.output_file.get()
        
        if not input_folder:
            messagebox.showerror("Error", "Please select a folder with papers.")
            return
        if not output_file:
            messagebox.showerror("Error", "Please specify an output file.")
            return
        if not os.path.isdir(input_folder):
            messagebox.showerror("Error", "Selected folder does not exist.")
            return
        
        try:
            verified_thresh = float(self.verified_threshold.get()) / 100
            partial_thresh = float(self.partial_threshold.get()) / 100
            if not (0 <= verified_thresh <= 1) or not (0 <= partial_thresh <= 1):
                raise ValueError("Thresholds must be between 0 and 100")
            if partial_thresh >= verified_thresh:
                raise ValueError("Partial threshold must be less than verified threshold")
        except ValueError as e:
            messagebox.showerror("Error", f"Invalid threshold values: {e}")
            return
        
        # Clear previous results
        self.clear_results()
        
        self.is_running = True
        self.run_btn.config(state=tk.DISABLED)
        self.progress.start()
        
        ignore_books = self.ignore_books.get()
        ignore_websites = self.ignore_websites.get()
        check_citations = not self.skip_citation_check.get()  # Inverted logic
        
        thread = threading.Thread(target=self.verification_worker,
                                  args=(input_folder, output_file, verified_thresh, partial_thresh, ignore_books, ignore_websites, check_citations))
        thread.start()
    
    def verification_worker(self, input_folder, output_file, verified_thresh, partial_thresh, ignore_books, ignore_websites, check_citations):
        try:
            self.update_status("Step 1/6: Reading papers...")
            lookup = ingest_papers(input_folder)
            if not lookup:
                self.root.after(0, lambda: messagebox.showerror("Error", "No papers found in folder."))
                return
            
            # Optional: Check citation-reference matching
            citation_results = None
            if check_citations:
                self.update_status("Step 2/7: Checking citation-reference matching...")
                citation_results = check_citation_matching(lookup)
            
            step_offset = 1 if check_citations else 0
            self.update_status(f"Step {2 + step_offset}/{6 + step_offset}: Extracting references...")
            results = extract_references_from_lookup(lookup)
            
            self.update_status(f"Step {3 + step_offset}/{6 + step_offset}: Splitting references...")
            split_results = split_references_from_results(results)
            
            self.update_status(f"Step {4 + step_offset}/{6 + step_offset}: Parsing references...")
            parsed = parse_all_references(split_results)
            
            self.update_status(f"Step {5 + step_offset}/{6 + step_offset}: Verifying references (this may take a while)...")
            verified = verify_all_references(parsed, delay=0.3,
                                             verified_threshold=verified_thresh,
                                             partial_threshold=partial_thresh,
                                             ignore_books=ignore_books,
                                             ignore_websites=ignore_websites)
            
            self.update_status(f"Step {6 + step_offset}/{6 + step_offset}: Generating report...")
            stats = generate_report(verified, output_file,
                                   verified_threshold=int(verified_thresh * 100),
                                   partial_threshold=int(partial_thresh * 100),
                                   ignore_books=ignore_books,
                                   ignore_websites=ignore_websites,
                                   citation_results=citation_results)
            
            # Capture values before scheduling callback
            final_stats = stats
            final_output = output_file
            final_ignore_books = ignore_books
            final_check_citations = check_citations
            
            self.root.after(0, lambda: self.show_results(final_stats, final_output, final_ignore_books, final_check_citations))
            self.root.after(0, lambda: messagebox.showinfo("Complete", f"Verification complete!\n\nReport saved to:\n{final_output}"))
        except Exception as e:
            import traceback
            error_msg = f"An error occurred:\n\n{str(e)}\n\n{traceback.format_exc()}"
            self.root.after(0, lambda: messagebox.showerror("Error", error_msg))
        finally:
            self.root.after(0, self.verification_complete)
    
    def verification_complete(self):
        self.is_running = False
        self.run_btn.config(state=tk.NORMAL)
        self.progress.stop()
        self.update_status("Verification complete!")
    
    def show_results(self, stats, output_file, ignore_books=False, check_citations=False):
        """Display results summary - removed to avoid duplication with report."""
        pass


def main():
    root = tk.Tk()
    app = ReferenceCheckerGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
